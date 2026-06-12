"""
Document Parser Tool for Xennic AI Platform

Parses various document formats:
- PDF (text, tables, metadata, images)
- DOCX (text, tables)
- TXT (plain text)
- Images (OCR with Tesseract)
"""

import os
import io
import tempfile
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field
import asyncio
from concurrent.futures import ThreadPoolExecutor

# PDF parsing
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    print("⚠️ PyMuPDF not installed. PDF parsing disabled.")

# DOCX parsing
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx not installed. DOCX parsing disabled.")

# OCR
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    print("⚠️ pytesseract or PIL not installed. OCR disabled.")


@dataclass
class ParsedDocument:
    """Result of document parsing"""
    file_type: str
    file_name: str
    text: str = ""
    tables: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: int = 0
    images: List[str] = field(default_factory=list)
    error: Optional[str] = None


class DocumentParser:
    """
    Parse various document formats
    
    Supported:
    - PDF (text, tables, images, metadata)
    - DOCX (text, tables)
    - TXT (plain text)
    - Images (OCR - English + Persian)
    """
    
    def __init__(self):
        self._executor = ThreadPoolExecutor(max_workers=2)
    
    async def parse(self, file_path: str, file_name: str) -> ParsedDocument:
        """
        Parse document based on file extension
        
        Args:
            file_path: Path to the file
            file_name: Original file name (for extension detection)
            
        Returns:
            ParsedDocument with extracted content
        """
        ext = os.path.splitext(file_name)[1].lower()
        
        if ext == '.pdf':
            return await self._parse_pdf(file_path, file_name)
        elif ext in ['.docx', '.doc']:
            return await self._parse_docx(file_path, file_name)
        elif ext == '.txt':
            return await self._parse_txt(file_path, file_name)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return await self._parse_image(file_path, file_name)
        else:
            return ParsedDocument(
                file_type="unknown",
                file_name=file_name,
                error=f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt, .png, .jpg, .jpeg",
            )
    
    async def _parse_txt(self, file_path: str, file_name: str) -> ParsedDocument:
        """Parse plain text file"""
        def _sync_parse():
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return ParsedDocument(
                file_type="txt",
                file_name=file_name,
                text=text,
                pages=1,
                metadata={"encoding": "utf-8"},
            )
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self._executor, _sync_parse)
    
    async def _parse_pdf(self, file_path: str, file_name: str) -> ParsedDocument:
        """Parse PDF file"""
        if not HAS_FITZ:
            return ParsedDocument(
                file_type="pdf",
                file_name=file_name,
                error="PyMuPDF not installed",
            )
        
        def _sync_parse():
            doc = fitz.open(file_path)
            
            result = ParsedDocument(
                file_type="pdf",
                file_name=file_name,
                pages=doc.page_count,
                metadata=dict(doc.metadata) if doc.metadata else {},
            )
            
            # Extract text from all pages
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_parts.append(page.get_text())
                
                # Extract tables (simplified)
                try:
                    tables = page.find_tables()
                    for table in tables:
                        table_data = table.extract()
                        if table_data:
                            result.tables.append({
                                "page": page_num + 1,
                                "rows": len(table_data),
                                "cols": len(table_data[0]) if table_data else 0,
                                "data": table_data,
                            })
                except Exception:
                    pass
            
            result.text = "\n\n".join(text_parts)
            doc.close()
            return result
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self._executor, _sync_parse)
    
    async def _parse_docx(self, file_path: str, file_name: str) -> ParsedDocument:
        """Parse DOCX file"""
        if not HAS_DOCX:
            return ParsedDocument(
                file_type="docx",
                file_name=file_name,
                error="python-docx not installed",
            )
        
        def _sync_parse():
            doc = Document(file_path)
            
            result = ParsedDocument(
                file_type="docx",
                file_name=file_name,
            )
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            result.text = "\n\n".join(text_parts)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    result.tables.append({
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                        "data": table_data,
                    })
            
            # Extract core properties (metadata)
            core_props = doc.core_properties
            if core_props:
                result.metadata = {
                    "author": core_props.author,
                    "title": core_props.title,
                    "subject": core_props.subject,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                }
            
            return result
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self._executor, _sync_parse)
    
    async def _parse_image(self, file_path: str, file_name: str) -> ParsedDocument:
        """Parse image with OCR"""
        if not HAS_OCR:
            return ParsedDocument(
                file_type="image",
                file_name=file_name,
                error="OCR libraries not installed",
            )
        
        def _sync_parse():
            image = Image.open(file_path)
            
            result = ParsedDocument(
                file_type="image",
                file_name=file_name,
                metadata={
                    "width": image.width,
                    "height": image.height,
                    "mode": image.mode,
                },
            )
            
            # OCR with English + Persian support
            try:
                text = pytesseract.image_to_string(image, lang='eng+fas')
                result.text = text.strip()
            except Exception as e:
                # Fallback to English only
                try:
                    text = pytesseract.image_to_string(image, lang='eng')
                    result.text = text.strip()
                except Exception as e2:
                    result.error = f"OCR failed: {str(e2)}"
            
            return result
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self._executor, _sync_parse)
    
    async def extract_text_only(self, file_path: str, file_name: str) -> str:
        """Extract only text from document (no metadata/tables)"""
        result = await self.parse(file_path, file_name)
        return result.text if not result.error else ""
    
    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions"""
        extensions = ['.txt']
        if HAS_FITZ:
            extensions.append('.pdf')
        if HAS_DOCX:
            extensions.append('.docx')
        if HAS_OCR:
            extensions.extend(['.png', '.jpg', '.jpeg', '.tiff', '.bmp'])
        return extensions
